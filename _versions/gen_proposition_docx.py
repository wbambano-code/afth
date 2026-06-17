# -*- coding: utf-8 -*-
"""Genere la proposition financiere AFTh au format Word (.docx)."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# --- Palette AFTh ---
PRIMARY = RGBColor
BLEU = PRIMARY(0x1A, 0x5F, 0x7A)
EAU = PRIMARY(0x4F, 0xB8, 0xC9)
VIOLET = PRIMARY(0x6C, 0x3F, 0xA8)
GRIS_FOND = "E4EEF3"
BLANC = PRIMARY(0xFF, 0xFF, 0xFF)
NOIR = PRIMARY(0x22, 0x22, 0x22)

doc = Document()

# Police par defaut
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)
style.font.color.rgb = NOIR

def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)

def set_cell_text(cell, text, bold=False, color=None, align=None, size=10):
    cell.text = ''
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color

def h1(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = BLEU
    p.space_before = Pt(14)
    p.space_after = Pt(6)
    return p

def h2(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = BLEU
    p.space_before = Pt(12)
    p.space_after = Pt(4)
    return p

def para(text, italic=False, size=10.5, color=None, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    r.bold = bold
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    return p

def style_header_row(row, fill=None):
    fill = fill or '1A5F7A'
    for cell in row.cells:
        shade(cell, fill)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = BLANC
                run.bold = True

# ============ EN-TETE ============
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = title.add_run("Proposition financière — Refonte du site AFTh")
r.bold = True
r.font.size = Pt(20)
r.font.color.rgb = BLEU

sub = doc.add_paragraph()
r = sub.add_run("Détail poste à poste — préparé par WTC (William) · Association Française des Techniques Hydrothermales")
r.italic = True
r.font.size = Pt(10)
r.font.color.rgb = PRIMARY(0x66, 0x66, 0x66)

para("Document de travail destiné à la discussion budgétaire fonctionnalité par fonctionnalité. "
     "La colonne « prix proposé » correspond au tarif partenaire WTC ; la colonne « valeur marché » "
     "indique le coût d'une réalisation équivalente en agence classique (fourchette France pour une "
     "application sur-mesure de ce périmètre).", size=10)

# ============ 1. ONE-SHOT ============
h1("1 · Développement — coût one-shot")
para("Légende : 🔵 Cœur = inclus, structurant (non négociable)  ·  🟠 Retirable = la FTH peut décider de ne pas l'activer.", italic=True, size=9.5)

rows = [
    ("#", "Poste", "Ce qu'il couvre", "Statut", "Prix proposé", "Valeur marché"),
    ("1", "Socle technique & architecture", "Serveur, rendu dynamique, structure du site, mise en production, sécurité de base", "Cœur", "1 100 €", "~3 500 €"),
    ("2", "Récupération & migration des contenus", "224 liens, 119 articles, 22 bulletins, extraction images + 78 PDF, nettoyage XSS", "Cœur", "1 500 €", "~4 500 €"),
    ("3", "Modernisation graphique + logo", "Charte rafraîchie (sans refonte totale), rafraîchissement logo, responsive", "Cœur", "1 000 €", "~2 500 €"),
    ("4", "Site vitrine public", "Accueil, association, congrès, prix, contact (niveau 1)", "Cœur", "700 €", "~2 200 €"),
    ("5", "Actus + congrès + bulletins + bibliothèque articles", "Modules d'affichage et navigation des contenus", "Cœur", "900 €", "~2 800 €"),
    ("6", "Tags, filtres thématiques & recherche", "Indexation des 119 articles, filtres, moteur de recherche", "Cœur", "600 €", "~1 800 €"),
    ("7", "Espace membres (niveau 2)", "Comptes, connexion, profils, activation/reset email, import Excel des adhérents", "Cœur", "1 100 €", "~3 500 €"),
    ("8", "Back-office d'administration autonome", "Gestion actus/congrès/bulletins/membres/prompt IA sans dépendre du prestataire", "Cœur", "1 300 €", "~4 200 €"),
    ("9", "Recette, sécurité, documentation & formation", "Tests, audit sécurité, doc, prise en main du BO par l'équipe FTH", "Cœur", "600 €", "~2 000 €"),
    ("", "Sous-total cœur (non négociable)", "", "", "8 800 €", "~27 000 €"),
    ("10", "Forum + modération", "Catégories, fils, posts, outillage de modération", "Retirable", "700 €", "~2 500 €"),
    ("11", "Moteur IA conversationnel", "Interrogation IA des archives/bulletins (+ piste réglementation) — hors tokens", "Retirable", "500 €", "~3 000 €"),
    ("", "TOTAL one-shot", "", "", "10 000 €", "~32 500 €"),
]

t = doc.add_table(rows=len(rows), cols=6)
t.style = 'Table Grid'
t.alignment = WD_TABLE_ALIGNMENT.CENTER
widths = [Cm(0.8), Cm(4.2), Cm(6.0), Cm(1.8), Cm(2.2), Cm(2.4)]
for i, row in enumerate(rows):
    for j, val in enumerate(row):
        cell = t.rows[i].cells[j]
        align = WD_ALIGN_PARAGRAPH.RIGHT if j >= 4 else (WD_ALIGN_PARAGRAPH.CENTER if j in (0,3) else WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(cell, val, align=align, size=9.5)
        cell.width = widths[j]
# header
style_header_row(t.rows[0])
# sous-total et total : surlignage
for idx in (10, 13):
    for cell in t.rows[idx].cells:
        shade(cell, GRIS_FOND)
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
# total final en bleu
for cell in t.rows[13].cells:
    shade(cell, 'D6E6ED')

para("")
para("Argument clé : le back-office autonome (poste 8) est volontairement dans le cœur. C'est lui qui permet "
     "à la FTH de gérer son contenu sans coût récurrent — l'enlever ferait économiser ~1 300 € une fois, "
     "mais générerait des factures de modification à chaque actu ou bulletin.", bold=False, italic=True, size=10)

# ============ 2. RECURRENT ============
h1("2 · Récurrent — environ 2 000 €/an")
rows2 = [
    ("Poste", "Ce qu'il couvre", "Montant indicatif"),
    ("Hébergement", "VPS OVH FR (RGPD), SSL auto, nom de domaine", "~150 €/an"),
    ("Sauvegardes & supervision", "Backup chiffré hebdo (Object Storage), surveillance", "~150 €/an"),
    ("Maintenance corrective & sécurité", "Mises à jour Node/dépendances, correctifs, patchs", "~700 €/an"),
    ("Maintenance évolutive", "Petits ajustements structurels, demandes ponctuelles", "~500 €/an"),
    ("Bulletin annuel + MAJ congrès/actus", "Mise en ligne du bulletin + cycle congrès de saison", "~350 €/an"),
    ("Support FTH", "Assistance email, aide back-office", "~150 €/an"),
    ("TOTAL", "", "~2 000 €/an"),
]
t2 = doc.add_table(rows=len(rows2), cols=3)
t2.style = 'Table Grid'
t2.alignment = WD_TABLE_ALIGNMENT.CENTER
widths2 = [Cm(5.0), Cm(8.5), Cm(3.0)]
for i, row in enumerate(rows2):
    for j, val in enumerate(row):
        cell = t2.rows[i].cells[j]
        align = WD_ALIGN_PARAGRAPH.RIGHT if j == 2 else WD_ALIGN_PARAGRAPH.LEFT
        set_cell_text(cell, val, align=align, size=9.5)
        cell.width = widths2[j]
style_header_row(t2.rows[0])
for cell in t2.rows[-1].cells:
    shade(cell, GRIS_FOND)
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True
para("")
para("Si le forum est retiré → la charge de modération + maintenance associée disparaît (≈ −150 à −250 €/an).", italic=True, size=10)

# ============ 3. IA ============
h1("3 · IA — budget tokens (piloté par la FTH)")
for txt in [
    "Séparé du forfait, payé en direct par la FTH (CB FTH, plafond mensuel défini dans la console).",
    "Plafond proposé : 5 $/mois — réaliste pour 100 à 200 adhérents en usage léger, à condition d'utiliser un modèle économique (type Claude Haiku) et la mise en cache des archives.",
    "À surveiller : si l'usage augmente fortement (ou si l'on branche la réglementation technique/sanitaire, plus volumineuse), prévoir de pouvoir relever le plafond. Le pilotage reste 100 % côté FTH — pas de surprise de facturation via WTC.",
]:
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(txt)
    r.font.size = Pt(10.5)

# ============ 4. SCENARIOS ============
h1("4 · Scénarios d'arbitrage")
rows3 = [
    ("Scénario", "One-shot", "Récurrent", "IA tokens"),
    ("Complet (recommandé)", "10 000 €", "~2 000 €/an", "~5 $/mois"),
    ("Sans forum", "9 300 €", "~1 800 €/an", "~5 $/mois"),
    ("Sans IA", "9 500 €", "~2 000 €/an", "0 €"),
    ("Sans forum ni IA", "8 800 €", "~1 800 €/an", "0 €"),
]
t3 = doc.add_table(rows=len(rows3), cols=4)
t3.style = 'Table Grid'
t3.alignment = WD_TABLE_ALIGNMENT.CENTER
widths3 = [Cm(6.0), Cm(3.2), Cm(3.5), Cm(3.0)]
for i, row in enumerate(rows3):
    for j, val in enumerate(row):
        cell = t3.rows[i].cells[j]
        align = WD_ALIGN_PARAGRAPH.RIGHT if j >= 1 else WD_ALIGN_PARAGRAPH.LEFT
        set_cell_text(cell, val, align=align, size=9.5)
        cell.width = widths3[j]
style_header_row(t3.rows[0])
# ligne recommandee surlignee
for cell in t3.rows[1].cells:
    shade(cell, 'D6E6ED')
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True
para("")
para("Recommandation de présentation : présenter le complet à 10 000 € comme l'offre de référence "
     "(tout est déjà développé dans la maquette = pas de surcoût caché), puis les retraits forum/IA "
     "comme de simples leviers d'économie marginale (−700 € / −500 €). Couper ne fait quasiment pas "
     "baisser la facture, mais fait perdre les deux fonctions les plus différenciantes.", italic=True, size=10)

# Pied de page
para("")
foot = doc.add_paragraph()
r = foot.add_run("Document indicatif — WTC · contact AFTh : 1 rue Cels, 75014 Paris · info@afth.asso.fr · 01 53 91 05 75")
r.italic = True
r.font.size = Pt(8.5)
r.font.color.rgb = PRIMARY(0x88, 0x88, 0x88)

out = r"C:\afth\Proposition-financiere-AFTh.docx"
doc.save(out)
print("OK ->", out)
